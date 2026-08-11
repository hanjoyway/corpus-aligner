"""Export aligned sentence pairs to translator-friendly formats.

Each exporter takes pairs ([{"zh","en","type","score"}]) and returns
(bytes, mimetype, filename). Content is never altered — pure formatting.
"""
from __future__ import annotations

import html
import io


def _esc(s: str) -> str:
    return html.escape(s or "", quote=False)


def to_tmx(pairs, base="aligned", src="zh-CN", tgt="en-US"):
    tus = []
    for p in pairs:
        tus.append(
            f'    <tu>\n'
            f'      <tuv xml:lang="{src}"><seg>{_esc(p["zh"])}</seg></tuv>\n'
            f'      <tuv xml:lang="{tgt}"><seg>{_esc(p["en"])}</seg></tuv>\n'
            f'    </tu>'
        )
    xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n<tmx version="1.4">\n'
        '  <header creationtool="平行语料对齐工作台" creationtoolversion="1.0" '
        f'segtype="sentence" o-tmf="plain" adminlang="en" srclang="{src}" datatype="plaintext"/>\n'
        '  <body>\n' + "\n".join(tus) + "\n  </body>\n</tmx>\n"
    )
    return xml.encode("utf-8"), "application/x-tmx+xml", f"{base}.tmx"


def to_xlsx(pairs, base="aligned"):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "对齐"
    ws.append(["序号", "中文", "English", "相似度"])
    for i, p in enumerate(pairs, 1):
        ws.append([i, p.get("zh", ""), p.get("en", ""), p.get("score", "")])
    ws.column_dimensions["B"].width = 55
    ws.column_dimensions["C"].width = 65
    for row in ws.iter_rows():
        for c in row:
            c.alignment = openpyxl.styles.Alignment(wrap_text=True, vertical="top")
    buf = io.BytesIO(); wb.save(buf)
    return buf.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", f"{base}.xlsx"


def to_docx(pairs, base="aligned"):
    import docx
    d = docx.Document()
    for p in pairs:
        d.add_paragraph(p.get("zh", ""))
        para = d.add_paragraph()
        run = para.add_run(p.get("en", ""))
        run.italic = True
        d.add_paragraph("")
    buf = io.BytesIO(); d.save(buf)
    return buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{base}.docx"


def to_tsv(pairs, base="aligned"):
    lines = ["zh\ten"] + [f'{p.get("zh","")}\t{p.get("en","")}' for p in pairs]
    return "\n".join(lines).encode("utf-8"), "text/tab-separated-values", f"{base}_lr.txt"


def to_txt_stacked(pairs, base="aligned"):
    blocks = [f'{p.get("zh","")}\n{p.get("en","")}' for p in pairs]
    return "\n\n".join(blocks).encode("utf-8"), "text/plain", f"{base}_stacked.txt"


def to_bfsu(pairs, base="aligned"):
    """Two line-aligned txt files (ZH.txt / EN.txt), one segment per line —
    the format BFSU ParaConc4Mac expects. Packaged as a zip."""
    import zipfile
    zh_lines, en_lines = [], []
    for p in pairs:
        z = (p.get("zh", "") or "").replace("\n", " ").strip()
        e = (p.get("en", "") or "").replace("\n", " ").strip()
        if not z or not e:
            continue  # keep only complete pairs so the two files stay line-aligned
        zh_lines.append(z)
        en_lines.append(e)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr(f"{base}.ZH.txt", "\n".join(zh_lines))
        z.writestr(f"{base}.EN.txt", "\n".join(en_lines))
    return buf.getvalue(), "application/zip", f"{base}_bfsu.zip"


# ---- ABBYY 风格双语 RTF（治国理政建库交付要求的 N_c.rtf）----
# 格式逐字对照 ABBYY Aligner 2.0「Export to Bilingual RTF」真实输出：单表双列，
# 每个对齐句对一行（左中文 / 右英文），CJK 用 \uN\'3f 转义。
_RTF_HEAD = ("{\\rtf1 \\deff1{\\fonttbl {\\f1\\fcharset134\\cpg936 SimSun;}\n}\n"
             "{\\colortbl }\n{\\stylesheet }\n{\\*\\listtable }\n{\\*\\listoverridetable }\n"
             "\\lnbrkrule\n\\sectd\n")
_RTF_CELLS = ("\\trgaph100\\trpaddl100\\trpaddfl3\\trpaddr100\\trpaddfr3"
              "\\clbrdrt\\brdrs\\brdrw10\\clbrdrb\\brdrs\\brdrw10\\clbrdrl\\brdrs\\brdrw10"
              "\\clbrdrr\\brdrs\\brdrw10\\cellx4731 "
              "\\clbrdrt\\brdrs\\brdrw10\\clbrdrb\\brdrs\\brdrw10\\clbrdrl\\brdrs\\brdrw10"
              "\\clbrdrr\\brdrs\\brdrw10\\cellx9462 ")


def _rtf_esc(s: str) -> str:
    """RTF 文本转义：控制字符转义，非 ASCII 转 \\uN\\'3f（N 为带符号 16 位）。"""
    out = []
    for ch in (s or "").replace("\n", " ").replace("\r", " "):
        o = ord(ch)
        if ch in "\\{}":
            out.append("\\" + ch)
        elif o < 128:
            out.append(ch)
        else:
            out.append(f"\\u{o - 65536 if o > 32767 else o}\\'3f")
    return "".join(out)


def to_rtf(pairs, base="aligned"):
    rows = []
    for p in pairs:
        z, e = _rtf_esc(p.get("zh", "")), _rtf_esc(p.get("en", ""))
        rows.append(
            f"\\trowd {_RTF_CELLS}\\intbl\\itap1\n"
            f"\\pard\\plain\\intbl\\itap1\\plain \\fs24 {z}\\cell \n"
            f"\\pard\\plain\\intbl\\itap1\\plain \\fs24 {e}\\cell \n"
            f"\\trowd {_RTF_CELLS}\\row \n"
        )
    rtf = _RTF_HEAD + "".join(rows) + "}\n"
    return rtf.encode("utf-8"), "application/rtf", f"{base}_双语对照.rtf"


def to_jianku(pairs, base="aligned"):
    """治国理政建库交付四件套打包。N_c.txt / N_e.txt / N_c.rtf 直接生成；
    N_c.ata 为 ABBYY 私有工程格式（见 README）——本工具不冒充生成，留给人工补。"""
    import zipfile
    zh_lines, en_lines = [], []
    for p in pairs:
        z = (p.get("zh", "") or "").replace("\n", " ").strip()
        e = (p.get("en", "") or "").replace("\n", " ").strip()
        if not z or not e:
            continue
        zh_lines.append(z)
        en_lines.append(e)
    rtf_bytes, _, _ = to_rtf(pairs, base)
    readme = (
        f"治国理政建库交付 · 文本编号 {base}\n"
        "本工具自动生成 3/4 文件：\n"
        f"  {base}_c.txt   中文，逐句一行（UTF-8）\n"
        f"  {base}_e.txt   英文，逐句一行（UTF-8）\n"
        f"  {base}_c.rtf   双语对照 RTF（格式对齐 ABBYY「Bilingual RTF」）\n"
        f"缺 {base}_c.ata（ABBYY Aligner 私有工程格式）——需在 ABBYY 中打开本 RTF/txt 另存，\n"
        "或与建库方确认是否必须提交 .ata。\n"
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr(f"{base}/{base}_c.txt", "\n".join(zh_lines))
        z.writestr(f"{base}/{base}_e.txt", "\n".join(en_lines))
        z.writestr(f"{base}/{base}_c.rtf", rtf_bytes)
        z.writestr(f"{base}/README.txt", readme)
    return buf.getvalue(), "application/zip", f"{base}_建库交付.zip"


def to_all(pairs, base="aligned"):
    """一个 zip 打包全部格式:TMX / Excel / 上下对照 Word / 左右对照 txt /
    上下对照 txt / 双语对照 RTF / BFSU 双 txt(逐行 ZH+EN)。通用「导出全部」。"""
    import zipfile
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for fn in (to_tmx, to_xlsx, to_docx, to_tsv, to_txt_stacked, to_rtf):
            content, _mime, name = fn(pairs, base)
            z.writestr(name, content)
        bz, _m, _n = to_bfsu(pairs, base)          # bfsu 本身是 zip,解包平铺进来
        inner = zipfile.ZipFile(io.BytesIO(bz))
        for n in inner.namelist():
            z.writestr(n, inner.read(n))
    return buf.getvalue(), "application/zip", f"{base}_全部格式.zip"


EXPORTERS = {"tmx": to_tmx, "xlsx": to_xlsx, "docx": to_docx,
             "tsv": to_tsv, "txt": to_txt_stacked, "bfsu": to_bfsu,
             "rtf": to_rtf, "jianku": to_jianku, "all": to_all}
